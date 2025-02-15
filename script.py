from docx import Document
from docx.shared import Pt
import settings

def add_word(new):
    document = Document(f'./dictionaries/{settings.dictionary}')

    core_properties = document.core_properties
    core_properties.author = "Dictionary"

    last_modified_by = document.core_properties
    last_modified_by.last_modified_by = "Dictionary"
    
    a = []
    for i in document.paragraphs:
        if i != '':
            a.append(i.text)
    a.append(new)
    a.sort()

    for k in document.paragraphs:
        delete_paragraph(k)
    
    for j in a:
        par = document.add_paragraph()
        run = par.add_run(j)
        font = run.font
        font.size = Pt(18)

    document.save(f'./dictionaries/{settings.dictionary}')

def add_dictionary(name):
    Document().save(f"./dictionaries/{name}.docx")

# def rename_dictionary(name, new_name):
#     document = Document(f"./dictionaries/{name}")
#     document.save(f"./dictionaries/{new_name}")

def remove_i_paragraph(paragraph):
    document = Document(f'./dictionaries/{settings.dictionary}')

    core_properties = document.core_properties
    core_properties.author = "Dictionary"

    last_modified_by = document.core_properties
    last_modified_by.last_modified_by = "Dictionary"

    for i in document.paragraphs:
        if i.text == paragraph:
            delete_paragraph(i)
            break

    document.save(f'./dictionaries/{settings.dictionary}')

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def get_words():
    document = Document(f'./dictionaries/{settings.dictionary}')
    a = []
    for i in document.paragraphs:
        if i != '':
            a.append(i.text)
    return a